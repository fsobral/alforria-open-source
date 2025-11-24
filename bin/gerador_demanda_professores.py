import docx
import pandas as pd

def processa_professor(k, contrato, nome, lista, max_vagas, path=''):

    print('Processa professor ' + nome)

    d = docx.Document(path + 'modelo.docx')

    d.paragraphs[8].text = 'VAGA {0:d} DE {1:d}'.format(k, max_vagas)

    quadro1 = d.tables[2]

    quadro1.cell(1, 1).text = nome.upper()
    quadro1.cell(2, 1).text = contrato

    quadro2 = d.tables[3]

    chtt = 0

    last_row = quadro2.rows[-1]
    quadro2._tbl.remove(quadro2.rows[-1]._tr)

    for i, c in enumerate(lista):

        cod, turma, serie, disc, ch = c

        quadro2.cell(3 + i, 0).text = str(cod)
        quadro2.cell(3 + i, 1).text = str(turma)
        quadro2.cell(3 + i, 2).text = str(serie)
        quadro2.cell(3 + i, 3).text = str(disc)
        quadro2.cell(3 + i, 4).text = str(ch)

        chtt += ch

        if i >= 4:

            quadro2.add_row()
    
    last_row.cells[-1].text = str(chtt)
    quadro2._tbl.append(last_row._tr)

    d.save("{0:s}{1:02d} - {2:s}.docx".format(path, k, nome))


def cria_demandas(planilha, max_vagas, path=''):

    df = pd.read_csv(planilha)

    df['newcol'] = df.apply(lambda x: (x.Codigo, x.Turma, x.Serie, x.Disciplina, x.CH), axis=1)

    df2 = df.groupby(["Contrato", "Nome"]).agg({'newcol':lambda x: list(x)})

    for k, i in enumerate(df2.iterrows()):

        (contrato, nome) = i[0]

        lista = i[1].iloc[0]

        processa_professor(k + 1, contrato, nome, lista, max_vagas, path=path)
